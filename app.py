import streamlit as st
import re
import gc
import json
import io
import docx
import pdfplumber

from langchain_openai import ChatOpenAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# ==========================================
# PAGE CONFIGURATION & LAYOUT
# ==========================================
st.set_page_config(page_title="RFP Bid Response Evaluator", page_icon="📝", layout="wide")

# Custom CSS for Premium Design aesthetics
st.markdown("""
<style>
    .stApp {
        background-color: #f8f9fa;
        color: #212529;
        font-family: 'Inter', sans-serif;
    }
    .main .block-container {
        padding-top: 2rem;
    }
    h1, h2, h3 {
        color: #0056b3;
    }
    .metric-card {
        background: #ffffff;
        border-radius: 10px;
        padding: 20px;
        border: 1px solid #dee2e6;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
        margin-bottom: 20px;
        transition: transform 0.2s ease-in-out, box-shadow 0.2s ease-in-out;
    }
    .metric-card:hover {
        transform: translateY(-2px);
        border-color: #0056b3;
        box-shadow: 0 6px 12px rgba(0, 0, 0, 0.1);
    }
    .stButton>button {
        background: linear-gradient(135deg, #238636 0%, #2ea043 100%);
        color: white;
        border: none;
        border-radius: 6px;
        font-weight: 600;
        transition: background 0.3s;
    }
    .stButton>button:hover {
        background: linear-gradient(135deg, #2ea043 0%, #3fb950 100%);
    }
</style>
""", unsafe_allow_html=True)


# ==========================================
# RAG & DOCUMENT PROCESSING
# ==========================================
@st.cache_resource
def get_embeddings_model():
    """Load small local embeddings model."""
    return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

def extract_text_from_docx(file_bytes):
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = [para.text for para in doc.paragraphs]
    return "\n".join(full_text)

def extract_text_from_pdf(file_bytes):
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def process_and_embed_documents(uploaded_files, collection_name="reference_material"):
    """Parses, chunks, and embeds files into ChromaDB."""
    docs = []
    for uploaded_file in uploaded_files:
        file_bytes = uploaded_file.read()
        filename = uploaded_file.name.lower()
        text = ""
        if filename.endswith(".docx"):
            text = extract_text_from_docx(file_bytes)
        elif filename.endswith(".pdf"):
            text = extract_text_from_pdf(file_bytes)
        
        if text:
            docs.append(Document(page_content=text, metadata={"source": uploaded_file.name}))
            
    # Force GC after parsing
    gc.collect()

    if not docs:
        return None

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    splits = text_splitter.split_documents(docs)
    
    # Store in ChromaDB
    embeddings = get_embeddings_model()
    # Use ephemeral client in memory
    vectorstore = Chroma.from_documents(documents=splits, embedding=embeddings, collection_name=collection_name)
    
    gc.collect()
    return vectorstore

def extract_draft_from_full_pdf(file_bytes, target_question):
    """Chunks a full PDF and searches for the most relevant section to the target question."""
    text = extract_text_from_pdf(file_bytes)
    gc.collect()
    
    if not text:
        return ""
        
    doc = Document(page_content=text, metadata={"source": "full_bid_pack"})
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200) # larger chunks for draft context
    splits = text_splitter.split_documents([doc])
    
    embeddings = get_embeddings_model()
    vectorstore = Chroma.from_documents(documents=splits, embedding=embeddings, collection_name="temp_draft_extraction")
    
    results = vectorstore.similarity_search(target_question, k=1)
    
    # cleanup temp vectorstore
    vectorstore.delete_collection()
    gc.collect()
    
    if results:
        return results[0].page_content
    return ""


# ==========================================
# PYTHON-NATIVE LOGIC / HEURISTICS 
# ==========================================
def extract_themes_count(text, themes):
    """Pure Python keyword density checker for Win Themes."""
    counts = {}
    if not themes.strip():
        return counts
        
    theme_list = [t.strip().lower() for t in themes.split(",")]
    text_lower = text.lower()
    
    for theme in theme_list:
        matches = len(re.findall(r'\b' + re.escape(theme) + r'\b', text_lower))
        counts[theme] = matches
    return counts

def extract_shipley_metrics(text, client_name):
    """Pure Python lexical check for Shipley customer-focus (we vs client)."""
    text_lower = text.lower()
    inward_words = ["we", "us", "our", "ourselves"]
    
    inward_count = sum(len(re.findall(r'\b' + word + r'\b', text_lower)) for word in inward_words)
    client_name_count = len(re.findall(r'\b' + re.escape(client_name.strip().lower()) + r'\b', text_lower)) if client_name else 0
    
    return {
        "inward_pronoun_count": inward_count,
        "client_name_count": client_name_count,
        "ratio_client_to_inward": round(client_name_count / max(1, inward_count), 2)
    }

# ==========================================
# LLM AGENT INVOCATIONS
# ==========================================
def invoke_agent_with_fallback(llm, prompt_text, system_message="You are an expert RFP Evaluator."):
    """Invokes LLM and aggressively handles memory and error checking."""
    try:
        messages = [
            ("system", system_message),
            ("human", prompt_text)
        ]
        response = llm.invoke(messages)
        # Force garbage collection to free memory
        gc.collect()
        content = response.content.strip() if response.content else ""
        if not content:
            print("[DEBUG] Model returned an empty string for prompt:\n", prompt_text)
            return "⚠️ The model returned an empty response. This occasionally happens with smaller local models under strict constraints. Please try clicking 'Evaluate' again."
        return content
    except Exception as e:
        gc.collect()
        return f"Error evaluating agent: {e}"

def run_compliance_agent(llm, question, draft):
    prompt = f"""
Analyze if the draft response meets ALL functional requirements asked in the RFP question.
Be extremely brief. Output must end with a definite: PASS, FAIL, or PARTIAL.

RFP Question: '{question}'
Draft Response: '{draft}'

Provide a 2 sentence explanation, then end with: "STATUS: [PASS/FAIL/PARTIAL]"
"""
    return invoke_agent_with_fallback(llm, prompt)

def run_theme_agent(llm, draft, python_theme_stats, rag_context=""):
    context_injection = f"\n\nBest Practice Context:\n{rag_context}" if rag_context else ""
    prompt = f"""
We have tracked the following occurrences of Win Themes in the text using basic matching:
{python_theme_stats}

Given this draft response, evaluate conceptually if the themes are actually effectively woven into the narrative rather than just name-dropped.
Draft Response: '{draft}'{context_injection}

Provide a very short 1 paragraph explanation, and give a "Theme Quality Score" percentage (e.g., 85%).
"""
    return invoke_agent_with_fallback(llm, prompt)

def run_shipley_agent(llm, draft, python_shipley_stats, rag_context=""):
    context_injection = f"\n\nBest Practice Context:\n{rag_context}" if rag_context else ""
    prompt = f"""
Shipley metrics calculated:
- First-person inward words (we/us/our): {python_shipley_stats['inward_pronoun_count']}
- Client Name Mentions: {python_shipley_stats['client_name_count']}

Evaluate the draft for active voice, clarity, and structural Shipley customer-focus.
Draft text: '{draft}'{context_injection}

Provide a brief paragraph on Active voice and Clarity, followed by a final "Shipley Score" out of 10.
"""
    return invoke_agent_with_fallback(llm, prompt)

# ==========================================
# UI BUILD & ORCHESTRATION 
# ==========================================
def main():
    st.title("🚀 RFP Bid Response Evaluator")
    st.markdown("Phase 2: RAG-Augmented, Local LLM-as-a-Judge using Hybrid Python + GenAI logic.")

    # Initialize Session State
    if "ref_vectorstore" not in st.session_state:
        st.session_state.ref_vectorstore = None
    if "extracted_draft" not in st.session_state:
        st.session_state.extracted_draft = ""

    with st.sidebar:
        st.header("⚙️ Local Server Config")
        local_url = st.text_input("LM Studio Endpoint", "http://localhost:1234/v1")
        model_name = st.text_input("Model Name", "Nemotron-3-nano-4b")
        st.caption("Settings bounded for speed and low CPU utilization.")
        
        st.divider()
        st.header("📋 Evaluation Context")
        win_themes = st.text_area("Win Themes (Comma Separated)", "Cost Efficiency, Seamless Integration, 24/7 Support")
        client_name = st.text_input("Target Client Name", "Acme Corp")
        
        st.divider()
        st.header("📚 Reference Knowledge Base")
        ref_files = st.file_uploader("Upload Reference Material (PDF/Word)", type=["pdf", "docx"], accept_multiple_files=True)
        if st.button("Process Knowledge Base"):
            if ref_files:
                with st.spinner("Processing and Embedding documents..."):
                    st.session_state.ref_vectorstore = process_and_embed_documents(ref_files)
                st.success("Knowledge Base updated!")
            else:
                st.warning("Please upload files first.")
        
        if st.session_state.ref_vectorstore:
            st.success("✅ Knowledge Base Active")

    st.subheader("Input Method")
    input_mode = st.radio("Select Input Mode:", ["Targeted Input (Word Doc)", "Full Bid Pack (PDF)"], horizontal=True)

    rfp_q = ""
    rfp_draft = ""

    if input_mode == "Targeted Input (Word Doc)":
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("RFP Question / Requirement")
            rfp_q = st.text_area("Paste the exact requirement here...", height=150)
            
        with col2:
            st.subheader("Draft Bid Response (.docx)")
            draft_file = st.file_uploader("Upload drafted answer", type=["docx"])
            if draft_file:
                rfp_draft = extract_text_from_docx(draft_file.read())
                st.success("Document parsed successfully.")
                with st.expander("Preview Draft Text"):
                    st.write(rfp_draft[:500] + ("..." if len(rfp_draft) > 500 else ""))

    else:
        st.subheader("Full Bid Pack (PDF)")
        pdf_file = st.file_uploader("Upload full RFP / Bid Pack", type=["pdf"])
        rfp_q = st.text_area("Target Question (to extract draft response)", height=100)
        
        if pdf_file and rfp_q:
            if st.button("Extract Draft Response from PDF"):
                with st.spinner("Searching PDF for relevant draft section..."):
                    extracted_text = extract_draft_from_full_pdf(pdf_file.read(), rfp_q)
                    st.session_state.extracted_draft = extracted_text
                if extracted_text:
                    st.success("Draft extracted successfully.")
                else:
                    st.warning("Could not extract a relevant section.")

        if st.session_state.extracted_draft:
            rfp_draft = st.session_state.extracted_draft
            with st.expander("Preview Extracted Draft Text", expanded=True):
                st.write(rfp_draft)

    st.markdown("<br/>", unsafe_allow_html=True)
    if st.button("Evaluate Response with AI Pipeline (Sequential)", use_container_width=True):
        if not rfp_q or not rfp_draft:
            st.error("Please provide both the RFP Question and ensure a Draft Response is uploaded/extracted before evaluating.")
            return
            
        # Initialize LangChain LLM optimized for speed / max_tokens
        llm = ChatOpenAI(
            base_url=local_url,
            api_key="lm-studio",
            model=model_name,
            temperature=0.1,
            max_tokens=1500, # Significantly increased to accommodate reasoning tokens of the model
            timeout=180, # Generous timeout for local inference
            max_retries=1
        )
        
        st.markdown("<hr/>", unsafe_allow_html=True)
        st.subheader("📊 Evaluation Dashboard")
        
        # RAG Context Retrieval
        rag_context = ""
        if st.session_state.ref_vectorstore:
            with st.status("Retrieving Best Practices from Knowledge Base...", expanded=True) as status:
                retriever = st.session_state.ref_vectorstore.as_retriever(search_kwargs={"k": 2}) # Keep k=2 for low context window
                retrieved_docs = retriever.invoke(rfp_draft)
                rag_context = "\n".join([doc.page_content for doc in retrieved_docs])
                gc.collect()
                st.write(f"Retrieved {len(retrieved_docs)} chunks from ChromaDB.")
                status.update(label="RAG Context Retrieval completed.", state="complete", expanded=False)
        
        # Step 1: Python Engine Execution
        with st.status("Executing Python Lexical Analysis Engine...", expanded=True) as status:
            theme_metrics = extract_themes_count(rfp_draft, win_themes)
            shipley_metrics = extract_shipley_metrics(rfp_draft, client_name)
            st.write(f"Themes found lexically: {theme_metrics}")
            st.write(f"Shipley PR metrics: inward={shipley_metrics['inward_pronoun_count']}, client={shipley_metrics['client_name_count']}")
            status.update(label="Python Analytical execution completed successfully in < 1s.", state="complete", expanded=False)

        # Step 2: Sequential LLM Execution
        col_res1, col_res2, col_res3 = st.columns(3)
        
        # AGENT 1
        with col_res1:
            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
            st.subheader("🏛️ Agent 1: Compliance")
            with st.spinner("Analyzing Functional Fit..."):
                c_out = run_compliance_agent(llm, rfp_q, rfp_draft)
                st.write(c_out)
            st.markdown('</div>', unsafe_allow_html=True)

        # AGENT 2
        with col_res2:
            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
            st.subheader("🎯 Agent 2: Theme Qualifier")
            with st.spinner("Analyzing Narrative Win Themes..."):
                t_out = run_theme_agent(llm, rfp_draft, theme_metrics, rag_context)
                st.write(t_out)
            st.markdown('</div>', unsafe_allow_html=True)

        # AGENT 3
        with col_res3:
            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
            st.subheader("✍️ Agent 3: Shipley Grader")
            with st.spinner("Analyzing Tone and Focus..."):
                s_out = run_shipley_agent(llm, rfp_draft, shipley_metrics, rag_context)
                st.write(s_out)
            st.markdown('</div>', unsafe_allow_html=True)
            
        st.success("Sequential Multi-Agent Pipeline Completed. Memory successfully flushed.")

if __name__ == "__main__":
    main()
